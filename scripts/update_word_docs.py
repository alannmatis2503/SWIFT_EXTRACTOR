#!/usr/bin/env python3
"""
Met à jour les deux fichiers Word (DOCUMENTATION_TECHNIQUE et ARCHITECTURE)
à partir du contenu de livraison_dsi/GUIDE_TECHNIQUE.md v4.0.

Usage: python3 scripts/update_word_docs.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from copy import deepcopy
import os

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


# ─────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.find(qn('w:shd'))
    if shading_elm is None:
        from lxml import etree
        shading_elm = etree.SubElement(shading, qn('w:shd'))
    shading_elm.set(qn('w:fill'), color_hex)
    shading_elm.set(qn('w:val'), 'clear')


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, '2F5496')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


def add_heading(doc, text, level=1):
    """Add a heading with consistent formatting."""
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False, style=None):
    """Add a paragraph."""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10)
    return p


def add_bullet(doc, text, level=0):
    """Add a bullet point."""
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(text, style=style)
    for run in p.runs:
        run.font.size = Pt(10)
    return p


def add_code_block(doc, text):
    """Add a code block styled paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    # Light grey background via shading on paragraph
    return p


# ─────────────────────────────────────────────────
# DOCUMENTATION TECHNIQUE
# ─────────────────────────────────────────────────

def create_documentation_technique():
    """Regenerate DOCUMENTATION_TECHNIQUE_SWIFT_EXTRACTOR.docx."""
    doc = Document()

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('SWIFT EXTRACTOR\nDocumentation Technique')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(47, 84, 150)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Version 4.0 — Février 2026\nDate de dernière mise à jour : 23/02/2026')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # spacer

    # ── 1. Introduction ──
    add_heading(doc, '1. Introduction')
    add_para(doc, (
        "Ce document présente de manière exhaustive le fonctionnement de l'application PDF SWIFT Extractor. "
        "Il détaille pour chaque type de message (MT103, MT202, MT900, MT910, MT950) les champs extraits, "
        "les règles métier appliquées, les exceptions et la structure du fichier Excel généré."
    ))
    add_para(doc, (
        "L'application traite les fichiers PDF contenant des messages SWIFT et génère un classeur Excel "
        "multi-feuilles avec séparation des messages normaux, des exceptions et des doublons."
    ))

    # ── 2. Types de Messages SWIFT Traités ──
    add_heading(doc, '2. Types de Messages SWIFT Traités')
    add_heading(doc, '2.1 Messages Acceptés', level=2)
    add_para(doc, 'Les types de messages suivants sont acceptés et traités par l\'application :')
    add_table(doc,
        ['Type Interne', 'Code SWIFT', 'Description'],
        [
            ['fin.103', 'MT103', 'Virement client (single customer credit transfer)'],
            ['fin.202', 'MT202', 'Virement interbancaire (general financial institution transfer)'],
            ['fin.202.COV', 'MT202.COV', 'Virement interbancaire avec couverture'],
            ['fin.900', 'MT900', 'Confirmation de débit (mode analyse transferts uniquement)'],
            ['fin.910', 'MT910', 'Confirmation de crédit (credit confirmation)'],
            ['—', 'MT950', 'Relevé de compte (mode rapprochement MT950 uniquement)'],
        ])

    add_heading(doc, '2.2 Messages Rejetés', level=2)
    add_para(doc, (
        "Tout message SWIFT dont le type n'est pas 103, 202, 900, 910 ou 950 est automatiquement rejeté. "
        "Les messages MT210, MT199, etc. sont silencieusement ignorés. "
        "Le type 202.COV est accepté (le type de base 202 est extrait)."
    ))

    add_heading(doc, '2.3 Filtre NAK (sortants uniquement)', level=2)
    add_para(doc, (
        "En mode sortant, tout message dont les 600 premiers caractères contiennent le mot NAK "
        "(Negative Acknowledgement) est entièrement ignoré. Un message NAK est un message "
        "rejeté par le réseau SWIFT."
    ))

    # ── 3. Modes de fonctionnement ──
    add_heading(doc, '3. Modes de Fonctionnement')
    add_table(doc,
        ['Mode', 'Description', 'Entrée', 'Sortie'],
        [
            ['Messages entrants', 'Extraction des MT202, MT103, MT910 reçus', '1+ PDF', 'Excel multi-feuilles'],
            ['Messages sortants', 'Extraction des MT202, MT103 émis', '1+ PDF', 'Excel multi-feuilles'],
            ['Analyse transferts', 'Rapprochement MT900 ↔ MT103/MT202', '2× PDF', 'Excel multi-feuilles'],
            ['Rapprochement MT950', 'Rapprochement écritures F61 ↔ messages', '2× PDF', 'Excel multi-feuilles'],
        ])

    # ── 4. Colonnes du Classeur Excel ──
    add_heading(doc, '4. Colonnes du Classeur Excel et Sources d\'Extraction')
    add_heading(doc, '4.1 Liste des Colonnes', level=2)
    add_table(doc,
        ['Colonne', 'Description'],
        [
            ['code_banque', "Code BIC de l'institution"],
            ['date_reference', 'Date de valeur du transfert (DD/MM/YYYY)'],
            ['reference', 'Référence unique de la transaction (F20)'],
            ['reference_origine', "Référence d'origine (F21, MT202/MT910)"],
            ['type_MT', 'Type de message SWIFT (fin.202, fin.103, fin.910)'],
            ['pays_iso3', 'Code pays ISO 3 lettres (CMR, GAB…) ou BEAC'],
            ["Code du donneur d'ordre", "Code BIC ou numérique du donneur d'ordre"],
            ["donneur d'ordre", "Nom du donneur d'ordre"],
            ['Bénéficiaire', 'Nom du bénéficiaire'],
            ['correspondant', 'BIC du correspondant bancaire (sender ou receiver)'],
            ['montant', 'Montant numérique du transfert'],
            ['devise', 'Code devise ISO 3 lettres (XAF, USD, EUR…)'],
            ['commentaires', "Commentaires métier (F70/F72 ou motif d'exception)"],
            ['source_pdf', 'Nom du fichier PDF source (lien cliquable)'],
        ])

    # ── 4.2 MT103 Entrant ──
    add_heading(doc, '4.2 MT103 Entrant (Incoming)', level=2)
    add_table(doc,
        ['Colonne', 'Champ SWIFT Source', 'Détail'],
        [
            ['Code Banque', 'F52A', "IdentifierCode / Code d'identifiant"],
            ['Date', 'F32A', 'Date de valeur (format YYMMDD → ISO)'],
            ['Référence', 'F20', 'Transaction Reference Number'],
            ['Type MT', 'En-tête', 'Identifier: fin.103'],
            ['Pays', 'Mapping BIC', 'Obtenu depuis bic_codes.xlsx colonne Pays'],
            ["Donneur d'Ordre", 'F50K / F50F', 'BIC→nom, ou texte extrait du champ'],
            ['Bénéficiaire', '(Vide)', 'Non renseigné pour les MT103 entrants'],
            ['Correspondant', 'En-tête Sender', "BIC de l'expéditeur du message"],
            ['Montant', 'F32A', 'Section Amount/Montant'],
            ['Devise', 'F32A', 'Section Currency/Devise'],
            ['Commentaires', 'F70', 'Contenu complet du champ F70 (Remittance Info)'],
        ])

    add_para(doc, "Filtrage des faux BIC :", bold=True)
    add_para(doc, (
        "Les mots courants ressemblant à un BIC sont exclus de la détection : "
        "CAMEROON, GABON, FRANCE, CENTRAL, INTERNATIONAL, COMMERCIAL, NATIONAL, "
        "GENERALE, FINANCES, BANQUE, RECETTE, PAIERIE, TRESOR, MINISTERE, CAISSE, COMPTABLE, DETAILS."
    ))

    # ── 4.3 MT103 Sortant ──
    add_heading(doc, '4.3 MT103 Sortant (Outgoing)', level=2)
    add_para(doc, "Extraction du Donneur d'Ordre — Ordre de priorité :", bold=True)
    add_bullet(doc, 'Priorité 1 : Codes Trésor (1001–6001) dans F50F → colonne « Reglement » de bic_codes.xlsx')
    add_bullet(doc, 'Priorité 2 : Code BIC dans F52A → résolution via bic_codes.xlsx')
    add_bullet(doc, 'Priorité 3 : Codes CCF 4 chiffres dans F50F → colonne « CCF » de bic_codes.xlsx')
    add_bullet(doc, 'Priorité 4 : Nom après « Details: Détails: » dans F50F')
    add_bullet(doc, "Priorité 5 : Fallback — nom extrait par heuristique depuis F52A")

    add_table(doc,
        ['Colonne', 'Champ SWIFT Source', 'Détail'],
        [
            ['Code Banque', 'F52A', "IdentifierCode"],
            ['Date', 'F32A', 'Date de valeur'],
            ['Référence', 'F20', 'Transaction Reference Number'],
            ['Correspondant', 'En-tête Receiver', 'BIC du destinataire du message'],
            ['Montant / Devise', 'F32A', 'Amount et Currency'],
            ['Commentaires', 'F70', 'Contenu complet du champ F70'],
        ])

    # ── 4.4 MT202 ──
    add_heading(doc, '4.4 MT202 / MT202.COV Entrant (Incoming)', level=2)
    add_para(doc, "Extraction du donneur d'ordre :", bold=True)
    add_bullet(doc, 'Priorité 1 : Code BIC strict dans F52A (après IdentifierCode)')
    add_bullet(doc, 'Priorité 2 : Code 4 chiffres dans F52D (après PartyIdentifier) → colonne Reglement')
    add_bullet(doc, 'Priorité 3 : Code BIC dans F52D (excluant faux BIC)')
    add_bullet(doc, 'Priorité 4 : Nom textuel extrait directement de F52D')

    add_table(doc,
        ['Colonne', 'Champ SWIFT Source', 'Détail'],
        [
            ['Code Banque', 'F52A / F52D', 'Priorité: F52A (BIC), sinon F52D'],
            ['Date', 'F32A', 'ValueDate / Date de valeur'],
            ['Référence', 'F20', 'Transaction Reference'],
            ['Réf. Origine', 'F21', 'Related Reference'],
            ['Bénéficiaire', '(Vide)', 'Non renseigné pour les MT202 entrants'],
            ['Correspondant', 'En-tête Sender', "BIC de l'expéditeur"],
            ['Commentaires', 'F72', 'Texte descriptif (toutes occurrences, jointes par /)'],
        ])

    add_heading(doc, '4.5 MT202 / MT202.COV Sortant (Outgoing)', level=2)
    add_para(doc, "Extraction du bénéficiaire depuis F58A/F58D :", bold=True)
    add_bullet(doc, 'F58A (prioritaire) : extraction du code BIC → résolution via bic_codes.xlsx')
    add_bullet(doc, 'F58D (fallback) : BIC avant NameAndAddress, sinon texte brut')
    add_bullet(doc, 'Filtrage des faux BIC (mêmes mots exclus que pour MT103)')

    # ── 4.6 MT910 ──
    add_heading(doc, '4.6 MT910 Entrant (Incoming)', level=2)
    add_para(doc, (
        "Le MT910 est une confirmation de crédit. Le champ F52A représente à la fois "
        "le donneur d'ordre ET le bénéficiaire (même institution). "
        "Le pays (pays_iso3) est forcé via la résolution BIC (override du pays détecté par heuristique)."
    ))
    add_table(doc,
        ['Colonne', 'Champ SWIFT Source', 'Détail'],
        [
            ['Code Banque', 'F52A / F52D', 'Même logique que MT202'],
            ['Donneur / Bénéficiaire', 'F52A', 'Identiques (même institution)'],
            ['Correspondant', 'En-tête Sender', "BIC de l'expéditeur"],
            ['Commentaires', 'F72', 'Après « Info émetteur – destinataire », jusqu\'à Block 5'],
        ])

    # ── 4.7 MT900 ──
    add_heading(doc, '4.7 MT900 — Mode Analyse des Transferts', level=2)
    add_para(doc, (
        "Le MT900 est une confirmation de débit. Il est traité uniquement dans le mode « Analyse des transferts exécutés ». "
        "Le matching s'effectue en comparant la référence d'origine (F21) du MT900 avec la référence (F20) "
        "des messages MT103/MT202 (comparaison en majuscules)."
    ))

    # ── 4.8 MT950 ──
    add_heading(doc, '4.8 MT950 — Mode Rapprochement MT950', level=2)
    add_para(doc, (
        "Le MT950 est un relevé de compte. Ses écritures F61 sont parsées et rapprochées "
        "avec les messages SWIFT extraits. Détail complet en section 7."
    ))
    add_table(doc,
        ['Champ F61', 'Description', 'Mapping interne'],
        [
            ['ValueDate', 'Date valeur (YYMMDD)', 'value_date'],
            ['DebitCreditMark', 'C (crédit) ou D (débit)', 'cd'],
            ['Amount', 'Montant entre #...#', 'amount'],
            ['IdentificationCode', 'Type de message (202, 910…)', 'identification_code'],
            ['ReferenceForTheAccountOwner', 'Référence propriétaire', 'ref_owner'],
            ['ReferenceOfTheAccountServicingInstitution', 'Réf. institution (après //)', 'ref_serv'],
            ['SupplementaryDetails', 'Détails complémentaires', 'supplementary_details'],
        ])

    # ── 5. Règles d'Exclusion et Exceptions ──
    add_heading(doc, "5. Règles d'Exclusion et Exceptions")
    add_para(doc, (
        "Chaque message extrait est soumis à une série de tests séquentiels. "
        "Le premier test positif détermine le routage, sauf la règle salle des marchés IBAN "
        "qui est évaluée en dernier."
    ))

    # 5.1 BDF
    add_heading(doc, '5.1 Exception BANQUE DE FRANCE (MT103 USD)', level=2)
    add_table(doc,
        ['Paramètre', 'Valeur'],
        [
            ['Type', 'fin.103 uniquement'],
            ['Devise', 'USD uniquement'],
            ['Champs inspectés', 'F53A, F54A, F57A'],
            ['Patterns recherchés', 'BANQUE DE FRANCE ou FW021083459'],
            ['Action', 'Message routé vers onglet BANQUE DE FRANCE'],
        ])

    # 5.2 BEACCMCX091 MT202
    add_heading(doc, '5.2 Exception BEACCMCX091 — MT202 sortant', level=2)
    add_table(doc,
        ['Paramètre', 'Valeur'],
        [
            ['Type', 'fin.202 (y compris 202.COV)'],
            ['Direction', 'outgoing uniquement'],
            ['Champ inspecté', 'F52A'],
            ['Pattern', 'BEACCMCX091 dans F52A'],
            ['Action', 'Message routé vers onglet Autres_Exceptions'],
            ['Annulation possible', 'Oui — si receiver_bic = SCBLGB2LXXX ou CITIGB2LXXX'],
        ])

    # 5.3 BC salle des marchés
    add_heading(doc, '5.3 Exception BC — Salle des marchés (MT202 sortant)', level=2)
    add_table(doc,
        ['Paramètre', 'Valeur'],
        [
            ['Type', 'fin.202'],
            ['Direction', 'outgoing uniquement'],
            ['Champ inspecté', 'F21 (référence d\'origine)'],
            ['Pattern', 'BC dans la référence d\'origine'],
            ['Action', 'Commentaire « opération salle des marchés » + Autres_Exceptions'],
            ['Annulation possible', 'Oui — si receiver_bic = SCBLGB2LXXX ou CITIGB2LXXX'],
        ])

    # 5.4 323201
    add_heading(doc, '5.4 Exception 323201 (MT202 entrant)', level=2)
    add_table(doc,
        ['Paramètre', 'Valeur'],
        [
            ['Type', 'fin.202 (y compris 202.COV)'],
            ['Direction', 'incoming uniquement'],
            ['Champ inspecté', 'F58A (Beneficiary Institution)'],
            ['Pattern', '323201 dans F58A'],
            ['Action', 'Message routé vers onglet Exceptions_323201'],
        ])

    # 5.5 EUR
    add_heading(doc, '5.5 Exceptions EUR — T2PI / T2RM / T2PL', level=2)
    add_para(doc, 'Ces exceptions s\'appliquent uniquement si la devise est EUR :')
    add_table(doc,
        ['Sous-règle', 'Pattern dans F20', 'Direction', 'Commentaire forcé'],
        [
            ['T2PI', 'T2PI', 'Entrant et sortant', 'intérêts'],
            ['T2RM', 'T2RM', 'Entrant uniquement', 'remboursement'],
            ['T2PL', 'T2PL', 'Sortant uniquement', 'placement'],
        ])
    add_para(doc, 'Action : Commentaire forcé + message routé vers onglet Autres_Exceptions.')

    # 5.6 Nivellement
    add_heading(doc, '5.6 Exception Nivellement (MT910)', level=2)
    add_para(doc, 'Type : fin.910 uniquement.')
    add_table(doc,
        ['Direction', 'Sous-règle', 'Champ inspecté', 'Pattern'],
        [
            ['Entrant', 'Règle 1', 'Référence (F20)', 'NIVLT'],
            ['Entrant', 'Règle 2', 'F25P ou F25', '5175'],
            ['Sortant', 'Règle 1', 'F53B ou F53', '5175'],
            ['Sortant', 'Règle 2', 'F58A ou F58', '5175'],
        ])
    add_para(doc, 'Action : Commentaire « nivellement » + message routé vers Autres_Exceptions.')

    # 5.7 BEACCMCX091 MT910
    add_heading(doc, '5.7 Exception BEACCMCX091 — MT910', level=2)
    add_table(doc,
        ['Paramètre', 'Valeur'],
        [
            ['Type', 'fin.910'],
            ['Champs inspectés', 'F50A puis F52A'],
            ['Pattern', 'Code BIC BEACCMCX091 dans IdentifierCode'],
            ['Action', 'Message routé vers onglet BEACCMCX091'],
        ])

    # 5.8 Forex
    add_heading(doc, '5.8 Exception Forex (MT910 entrant)', level=2)
    add_table(doc,
        ['Paramètre', 'Valeur'],
        [
            ['Type', 'fin.910'],
            ['Direction', 'incoming uniquement'],
            ['Condition', 'Les 8 premiers caractères du code_donneur_dordre dans la liste forex'],
            ['Source codes forex', 'Feuille « forex » de bic_codes.xlsx (colonne A, ligne 2+)'],
            ['Action', 'Commentaire « forex » + message routé vers onglet forex'],
        ])

    # 5.9 Salle marchés IBAN
    add_heading(doc, '5.9 Exception Salle des marchés — IBAN (MT910 entrant)', level=2)
    add_table(doc,
        ['Paramètre', 'Valeur'],
        [
            ['Type', 'fin.910'],
            ['Direction', 'incoming uniquement'],
            ['Champ inspecté', 'F25P ou F25'],
            ['Pattern', 'IBAN FR7630001000640000005169558'],
            ['Action', 'Commentaire « opération salle des marchés » + Autres_Exceptions'],
            ['Priorité', 'DERNIÈRE — uniquement si aucune autre exception'],
        ])

    # 5.10 Annulations
    add_heading(doc, '5.10 Annulation des exceptions', level=2)
    add_para(doc, (
        "Certaines exceptions MT202 sortant peuvent être annulées si le message est destiné "
        "à Standard Chartered Bank (SCBLGB2LXXX) ou Citibank (CITIGB2LXXX) de Londres."
    ))
    add_table(doc,
        ['Exception annulée', 'Condition d\'annulation', 'Effet'],
        [
            ['BEACCMCX091 MT202', 'sender_bic commence par BEACCMCX ET receiver_bic = SCBLGB2LXXX ou CITIGB2LXXX', 'Message redevient normal → summary'],
            ['BC salle marchés', 'receiver_bic = SCBLGB2LXXX ou CITIGB2LXXX', 'Message redevient normal, commentaire supprimé'],
        ])

    # 5.11 Exceptions MT900
    add_heading(doc, '5.11 Exceptions MT900 (Analyse des transferts)', level=2)
    add_table(doc,
        ['Pattern dans F20', 'Pattern dans F21', 'Exception'],
        [
            ['T2PL', '—', 'T2PL (placement)'],
            ['—', 'NIVLT ou NIVELLEMENT', 'nivellement'],
        ])
    add_para(doc, 'Ces MT900 ne sont PAS matchés avec les MT103/MT202 et vont dans une feuille Exceptions.')

    # ── 6. Structure du Classeur Excel Généré ──
    add_heading(doc, '6. Structure du Classeur Excel Généré')

    add_heading(doc, '6.1 Mode standard (entrants / sortants)', level=2)
    add_table(doc,
        ['Onglet', 'Condition', 'Contenu'],
        [
            ['summary', 'Toujours', 'Messages normaux (excluant doublons secondaires)'],
            ['BEACCMCX091', 'Si non vide', 'Messages BEACCMCX091 (MT910)'],
            ['Exceptions_323201', 'Si non vide', 'Messages avec 323201 dans F58A'],
            ['Autres_Exceptions', 'Si non vide', 'EUR, nivellement, salle des marchés, BC, BEACCMCX091-MT202'],
            ['BANQUE DE FRANCE', 'Si non vide', 'MT103 USD BANQUE DE FRANCE'],
            ['forex', 'Si non vide', 'MT910 forex'],
            ['Doublons_potentiels', 'Si doublons', 'Paires MT910/MT202 même référence + montant'],
            ['Par pays (CMR, GAB…)', 'Si pays renseigné', 'Sous-ensemble par pays, trié par correspondant'],
            ['Par message (nom PDF)', 'Pour chaque msg', 'Feuille clé/valeur détaillée'],
        ])

    add_heading(doc, '6.2 Détection des doublons', level=2)
    add_para(doc, (
        "Un MT910 peut être un doublon d'un MT202 lorsque les deux portent la même référence (F20, normalisée "
        "en majuscules) et le même montant. La recherche s'effectue sur toutes les listes confondues."
    ))
    add_bullet(doc, "Le premier message reste dans sa liste, commentaire préfixé « potentiel doublon »")
    add_bullet(doc, "Le second est exclu du summary")
    add_bullet(doc, "Les deux apparaissent dans l'onglet Doublons_potentiels")

    add_heading(doc, '6.3 Hyperliens bidirectionnels', level=2)
    add_para(doc, (
        "Chaque cellule source_pdf du summary est un lien cliquable vers la feuille de détail. "
        "Chaque feuille de détail contient un lien « ⬅ Retour au summary » en première ligne."
    ))

    add_heading(doc, '6.4 Mode Analyse des Transferts', level=2)
    add_table(doc,
        ['Onglet', 'Contenu'],
        [
            ['Transferts_Executes', 'MT900 matchés avec infos complétées du MT103/MT202'],
            ['MT900_non_rapproches', 'MT900 sans correspondant trouvé'],
            ['Suspens', 'MT103/MT202 sans confirmation MT900'],
            ['Exceptions', 'MT900 en exception (T2PL, nivellement)'],
        ])

    add_heading(doc, '6.5 Mode Rapprochement MT950', level=2)
    add_table(doc,
        ['Onglet', 'Contenu'],
        [
            ['Rapproches', 'Écritures F61 matchées côte-à-côte avec les messages'],
            ['Msg_non_rapproches', 'Messages sans écritures F61 correspondantes'],
            ['F61_non_rapproches', 'Écritures F61 sans message correspondant'],
        ])

    # ── 7. Mode Rapprochement MT950 — Détail ──
    add_heading(doc, '7. Mode Rapprochement MT950 — Détail Complet')
    add_para(doc, (
        "Ce mode rapproche les écritures F61 d'un relevé de compte MT950 avec les messages SWIFT "
        "(MT202, MT103, MT910) extraits de fichiers PDF séparés."
    ))

    add_heading(doc, '7.1 Sous-modes', level=2)
    add_table(doc,
        ['Sous-mode', 'Écritures F61', 'Messages'],
        [
            ['Entrants', 'Type C (crédit)', 'Messages entrants (MT202/MT103/MT910)'],
            ['Sortants', 'Type D (débit)', 'Messages sortants (MT202/MT103)'],
        ])

    add_heading(doc, '7.2 Pool de messages', level=2)
    add_para(doc, (
        "Tous les messages extraits par extract_dispatch() sont combinés pour le matching, "
        "y compris ceux routés vers les onglets spéciaux (BEAC, forex, BDF, exceptions). "
        "Seuls les types 202, 103, 910 sont conservés."
    ))

    add_heading(doc, '7.3 Clé de matching — 4 critères cumulatifs', level=2)
    add_para(doc, 'Les 4 critères suivants doivent être tous satisfaits simultanément :')
    add_table(doc,
        ['#', 'Critère', 'Condition', 'Tolérance'],
        [
            ['1', 'Montant', '|F61.amount − message.montant|', '≤ 0,011'],
            ['2', 'Type / Code identification', 'F61.identification_code contenu dans message.type_MT', 'Match partiel (ex: 202 dans fin.202)'],
            ['3', 'Date valeur', 'F61.value_date (YYMMDD → 20YY-MM-DD) = message.date_reference[:10]', 'Exact'],
            ['4', 'Référence', 'Dépend du sous-mode (voir ci-dessous)', 'Voir ci-dessous'],
        ])

    add_heading(doc, '7.4 Critère de référence : Entrants vs Sortants', level=2)
    add_table(doc,
        ['Sous-mode', 'Référence F61', 'Référence message', 'Comparaison'],
        [
            ['Entrants', 'ref_serv (après //)', 'reference (F20)', 'Match exact'],
            ['Sortants', 'ref_owner', 'reference (F20)', 'Match préfixe (l\'un commence par l\'autre)'],
        ])
    add_para(doc, (
        "Match préfixe : la fonction _refs_match(a, b) retourne True si a == b ou "
        "b.startswith(a) ou a.startswith(b). Cela couvre les cas où le relevé tronque la référence."
    ))

    add_heading(doc, '7.5 Numérotation F61', level=2)
    add_para(doc, (
        "Chaque écriture F61 reçoit un numéro séquentiel (f61_index, 1-based) pour traçabilité. "
        "Ce numéro apparaît dans la colonne « N° F61 » de l'onglet Rapproches et F61_non_rapproches."
    ))

    # ── 8. Structure du fichier bic_codes.xlsx ──
    add_heading(doc, '8. Structure du Fichier bic_codes.xlsx')
    add_heading(doc, '8.1 Feuille principale', level=2)
    add_table(doc,
        ['Colonne', 'Obligatoire', 'Description', 'Exemple'],
        [
            ['Code BIC', '✅', 'Code SWIFT 8 ou 11 caractères', 'ECOCMLCX'],
            ['Noms', '✅', "Nom complet de l'institution", 'ECOBANK CAMEROUN S.A.'],
            ['Pays', '✅', 'Code pays ISO3', 'CMR'],
            ['Reglement', '❌', 'Code de règlement à 4 chiffres', '8101'],
            ['CCF', '❌', 'Code CCF long', '10.311301.0.1401.0.0.0.0.0'],
        ])

    add_heading(doc, '8.2 Feuille forex', level=2)
    add_para(doc, (
        "Colonne A (à partir de la ligne 2) : codes BIC 8 caractères identifiant les correspondants forex."
    ))

    add_heading(doc, '8.3 Emplacements recherchés', level=2)
    add_bullet(doc, 'Variable d\'environnement PDF_SWIFT_DATA_DIR')
    add_bullet(doc, 'Dossier ProgramData (Windows) / LOCALAPPDATA')
    add_bullet(doc, 'Dossier utilisateur ~/.pdf_swift_extractor/data/')
    add_bullet(doc, 'data/bic_codes.xlsx (relatif à l\'application)')

    # ── 9. Extraction des commentaires ──
    add_heading(doc, '9. Extraction des Commentaires')
    add_table(doc,
        ['Type MT', 'Champ source', 'Logique'],
        [
            ['fin.202', 'F72', 'Toutes les occurrences de « Texte descriptif: » jointes par /'],
            ['fin.103', 'F70', 'Contenu complet du champ F70, nettoyé'],
            ['fin.910', 'F72', 'Après « Info émetteur – destinataire » jusqu\'à Block 5, jointes par /'],
        ])
    add_para(doc, (
        "Règle de non-écrasement : si un commentaire a été positionné par une règle d'exception "
        "(opération salle des marchés, forex, nivellement, intérêts…), l'extraction F70/F72 ne l'écrase pas."
    ))

    # ── 10. Extraction du correspondant ──
    add_heading(doc, '10. Extraction du Correspondant')
    add_table(doc,
        ['Direction', 'Valeur du correspondant'],
        [
            ['incoming (entrant)', "sender_bic — BIC de l'expéditeur du message"],
            ['outgoing (sortant)', 'receiver_bic — BIC du destinataire du message'],
        ])
    add_para(doc, (
        "Cas particulier MT900 : lors du matching avec un MT103/MT202, le correspondant du MT900 "
        "est remplacé par celui du MT103/MT202 associé."
    ))

    # ── Save ──
    path = os.path.join(BASE, 'DOCUMENTATION_TECHNIQUE_SWIFT_EXTRACTOR.docx')
    doc.save(path)
    print(f'✅ DOCUMENTATION_TECHNIQUE créé : {path}')
    print(f'   → {len(doc.paragraphs)} paragraphes, {len(doc.tables)} tables')


# ─────────────────────────────────────────────────
# ARCHITECTURE SWIFT EXTRACTOR
# ─────────────────────────────────────────────────

def create_architecture():
    """Regenerate ARCHITECTURE_SWIFT_EXTRACTOR.docx."""
    doc = Document()

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Architecture Technique\nApplication SWIFT Extractor')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(47, 84, 150)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Version 4.0 — Février 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()

    # ── Table des matières ──
    add_heading(doc, 'Table des Matières')
    toc_items = [
        '1. Vue d\'ensemble',
        '2. Architecture en couches',
        '3. Flux de traitement',
        '4. Composants détaillés',
        '5. Stack technique',
        '6. Règles métier implémentées',
        '7. Options de déploiement',
        '8. Caractéristiques techniques',
        '9. Points d\'attention pour le SI',
        '10. Conclusion',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')

    # ── 1. Vue d'ensemble ──
    add_heading(doc, '1. Vue d\'ensemble')
    add_para(doc, (
        "L'application SWIFT Extractor est une solution web développée en Python pour "
        "l'extraction automatisée et l'analyse de messages SWIFT depuis des fichiers PDF. "
        "Elle fonctionne en 4 modes : messages entrants, sortants, analyse des transferts exécutés, "
        "et rapprochement MT950."
    ))

    add_heading(doc, 'Objectifs', level=2)
    add_bullet(doc, 'Automatiser l\'extraction de données depuis les exports PDF SWIFT')
    add_bullet(doc, 'Structurer les informations dans des fichiers Excel multi-feuilles')
    add_bullet(doc, 'Classifier automatiquement les messages selon des règles métier')
    add_bullet(doc, 'Faciliter l\'analyse et le rapprochement des opérations bancaires')
    add_bullet(doc, 'Rapprocher les écritures MT950 avec les messages SWIFT')

    add_heading(doc, 'Types de messages traités', level=2)
    add_table(doc,
        ['Type', 'Description', 'Modes'],
        [
            ['MT103', 'Virement client', 'Entrants, Sortants, Analyse, MT950'],
            ['MT202/COV', 'Virement interbancaire', 'Entrants, Sortants, Analyse, MT950'],
            ['MT910', 'Confirmation de crédit', 'Entrants, MT950'],
            ['MT900', 'Confirmation de débit', 'Analyse transferts uniquement'],
            ['MT950', 'Relevé de compte', 'Rapprochement MT950 uniquement'],
        ])

    # ── 2. Architecture en couches ──
    add_heading(doc, '2. Architecture en couches')
    add_para(doc, (
        "L'application suit une architecture en couches classique pour garantir "
        "la séparation des responsabilités et la maintenabilité du code."
    ))

    add_heading(doc, 'Diagramme d\'architecture', level=2)
    add_code_block(doc, (
        "┌──────────────────────────────────────────────────────┐\n"
        "│         COUCHE PRÉSENTATION — app.py                │\n"
        "│    (Interface Streamlit — port 8501)                 │\n"
        "│  Upload PDF │ Modes │ Filtre dates │ Download        │\n"
        "└──────────────────────┬───────────────────────────────┘\n"
        "                       │\n"
        "┌──────────────────────▼───────────────────────────────┐\n"
        "│         COUCHE MÉTIER — extractor_manager.py         │\n"
        "│  Dispatching │ Workbook Excel │ Matching MT900/MT950 │\n"
        "└──────────────────────┬───────────────────────────────┘\n"
        "                       │\n"
        "┌──────────────────────▼───────────────────────────────┐\n"
        "│         COUCHE EXTRACTION — extractors/              │\n"
        "│  mt_multi │ mt202 │ mt103 │ mt910 │ mt900 │ mt950   │\n"
        "└──────────────────────┬───────────────────────────────┘\n"
        "                       │\n"
        "┌──────────────────────▼───────────────────────────────┐\n"
        "│         COUCHE DONNÉES — data/bic_codes.xlsx         │\n"
        "│       (Référentiel BIC — annuaire des banques)       │\n"
        "└──────────────────────────────────────────────────────┘"
    ))

    # 2.1 Couche Présentation
    add_heading(doc, '2.1 Couche Présentation', level=2)
    add_para(doc, 'Fichier principal : app.py', bold=True)
    add_para(doc, 'Technologies : Streamlit, Interface web responsive en français')
    add_para(doc, 'Fonctionnalités :', bold=True)
    add_bullet(doc, 'Upload de fichiers PDF (drag & drop, multi-fichiers)')
    add_bullet(doc, 'Sélection du mode de traitement (entrants / sortants / analyse transferts / rapprochement MT950)')
    add_bullet(doc, 'Prévisualisation multi-onglets des résultats Excel')
    add_bullet(doc, 'Filtrage par plage de dates (modes standard)')
    add_bullet(doc, 'Téléchargement des fichiers générés')
    add_bullet(doc, 'Affichage des statistiques de traitement')
    add_bullet(doc, 'Bouton Clear All pour réinitialiser')

    # 2.2 Couche Métier
    add_heading(doc, '2.2 Couche Métier', level=2)
    add_para(doc, 'A. Dispatcher Principal — extractors/mt_multi.py (~1750 lignes)', bold=True)
    add_para(doc, 'Responsabilités :')
    add_bullet(doc, "Extraction du texte brut depuis les PDFs (PyMuPDF prioritaire, pdfplumber en fallback)")
    add_bullet(doc, 'Découpage du PDF en messages individuels via 6 heuristiques regex')
    add_bullet(doc, 'Détection automatique du type de message SWIFT')
    add_bullet(doc, "Routage vers l'extracteur spécialisé approprié")
    add_bullet(doc, "Application des règles métier et exceptions (voir documentation technique)")
    add_bullet(doc, "Enrichissement des données (résolution BIC → noms)")
    add_bullet(doc, "Retour de 7 listes classifiées")

    add_para(doc, 'B. Extracteurs Spécialisés', bold=True)
    add_table(doc,
        ['Fichier', 'Lignes', 'Rôle'],
        [
            ['mt202.py', '~400', 'Extraction MT202 + utilitaires de base (champs, montants, dates)'],
            ['mt103.py', '~350', 'Extraction MT103 (donneur F50F/F50K, codes Trésor/CCF)'],
            ['mt910.py', '~200', 'Extraction MT910 (donneur = bénéficiaire)'],
            ['mt900.py', '~150', 'Extraction MT900 (F21 = clé de matching)'],
            ['mt950.py', '~280', 'Parsing F61 + matching 4 critères'],
        ])

    add_para(doc, 'C. Générateur Excel — extractor_manager.py (~1774 lignes)', bold=True)
    add_bullet(doc, 'Génération de fichiers Excel multi-feuilles :')
    add_bullet(doc, 'Feuille principale (messages normaux)', level=1)
    add_bullet(doc, 'Exceptions BEACCMCX091', level=1)
    add_bullet(doc, 'Exceptions 323201 (MT202 entrants)', level=1)
    add_bullet(doc, 'Autres exceptions (BC, EUR, nivellement, salle marchés)', level=1)
    add_bullet(doc, 'BANQUE DE FRANCE, forex', level=1)
    add_bullet(doc, 'Feuilles par pays', level=1)
    add_bullet(doc, 'Doublons potentiels', level=1)
    add_bullet(doc, 'Feuilles par fichier source (mode debug)', level=1)
    add_bullet(doc, 'Mise en forme automatique (colonnes, dates, montants)')
    add_bullet(doc, 'Hyperliens bidirectionnels entre feuilles')

    # 2.3 Couche Traitement
    add_heading(doc, '2.3 Couche Traitement', level=2)
    add_para(doc, 'A. Mapping BIC — extractors/bic_utils.py', bold=True)
    add_bullet(doc, 'Résolution codes BIC (8-11 caractères) → Nom institution')
    add_bullet(doc, 'Codes Trésor (1001-6001) → Nom pays')
    add_bullet(doc, 'Codes CCF 4 chiffres → Nom institution')
    add_bullet(doc, 'Codes forex → Détection MT910 forex')
    add_bullet(doc, 'Cache LRU (chargement Excel unique)')
    add_bullet(doc, 'Source de données : data/bic_codes.xlsx')

    # 2.4 Couche Données
    add_heading(doc, '2.4 Couche Données', level=2)
    add_para(doc, 'Entrées :', bold=True)
    add_bullet(doc, 'PDFs SWIFT : Messages exportés depuis systèmes bancaires')
    add_bullet(doc, 'bic_codes.xlsx : Référentiel des codes BIC et institutions')
    add_para(doc, 'Sorties :', bold=True)
    add_bullet(doc, 'Fichiers Excel (.xlsx) : Résultats structurés multi-feuilles')
    add_bullet(doc, 'Logs : Traçabilité du traitement (logs/app.log)')

    # ── 3. Flux de traitement ──
    add_heading(doc, '3. Flux de traitement')

    add_heading(doc, 'Étapes détaillées', level=2)
    steps = [
        ("1. Upload PDF", "L'utilisateur upload un ou plusieurs fichiers PDF et sélectionne le mode"),
        ("2. Initialisation", "L'interface Streamlit appelle extract_dispatch() ou les fonctions MT950"),
        ("3. Extraction texte", "PyMuPDF extrait le texte brut (~50× plus rapide, fallback pdfplumber)"),
        ("4. Découpage", "Le texte est découpé en messages individuels via 6 heuristiques regex en cascade"),
        ("5. Détection type", 'Pour chaque message, détection du type via "Identifier: fin.XXX"'),
        ("6. Extraction", "Routage vers l'extracteur spécialisé (mt202/mt103/mt910/mt900)"),
        ("7. Structuration", "L'extracteur retourne un dictionnaire avec les champs structurés"),
        ("8. Résolution BIC", "Les codes BIC sont résolus en noms d'institutions et pays via bic_codes.xlsx"),
        ("9. Règles métier", "Application des règles d'exception (BEACCMCX091, BC, 323201, EUR, nivellement, forex, BDF, salle marchés)"),
        ("10. Classification", "Le message est classé (normal, exception, doublon)"),
        ("11. Agrégation", "Tous les messages traités sont agrégés dans les 7 listes"),
        ("12. Génération Excel", "Le générateur crée le fichier Excel multi-feuilles"),
        ("13. Mise en forme", "Application des styles, formats de colonnes, ajustements, hyperliens"),
        ("14. Retour UI", "Le fichier est retourné à l'interface (bytes en mémoire)"),
        ("15. Présentation", "Prévisualisation multi-onglets + bouton téléchargement"),
    ]
    for step_title, step_desc in steps:
        p = doc.add_paragraph()
        run = p.add_run(step_title + ' : ')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(step_desc)
        run.font.size = Pt(10)

    # ── 4. Composants détaillés ──
    add_heading(doc, '4. Composants détaillés')
    add_heading(doc, '4.1 Structure des fichiers', level=2)
    add_code_block(doc, (
        "livraison_dsi/\n"
        "├── app.py                    ← Interface Streamlit\n"
        "├── extractor_manager.py      ← Dispatching + Excel\n"
        "├── utils.py                  ← Logging\n"
        "├── requirements.txt          ← Dépendances Python\n"
        "├── Dockerfile                ← Conteneurisation\n"
        "├── data/\n"
        "│   └── bic_codes.xlsx        ← Référentiel BIC\n"
        "├── extractors/\n"
        "│   ├── __init__.py\n"
        "│   ├── bic_utils.py          ← Résolution BIC\n"
        "│   ├── mt103.py              ← Extracteur MT103\n"
        "│   ├── mt202.py              ← Extracteur MT202\n"
        "│   ├── mt900.py              ← Extracteur MT900\n"
        "│   ├── mt910.py              ← Extracteur MT910\n"
        "│   ├── mt950.py              ← Parseur MT950 + matching\n"
        "│   └── mt_multi.py           ← Dispatcher multi-messages\n"
        "├── GUIDE_TECHNIQUE.md\n"
        "└── GUIDE_UTILISATEUR.md"
    ))

    add_heading(doc, '4.2 Colonnes extraites', level=2)
    add_table(doc,
        ['Colonne', 'Description'],
        [
            ['code_banque', "Code BIC de l'institution"],
            ['date_reference', 'Date valeur (DD/MM/YYYY)'],
            ['reference', 'Référence transaction (F20)'],
            ['reference_origine', "Référence d'origine (F21)"],
            ['type_MT', 'Type SWIFT (fin.202, fin.103…)'],
            ['pays_iso3', 'Code pays ISO3 ou BEAC'],
            ["Code du donneur d'ordre", "Code BIC du donneur d'ordre"],
            ["donneur d'ordre", "Nom du donneur d'ordre"],
            ['Bénéficiaire', 'Nom du bénéficiaire'],
            ['correspondant', 'BIC correspondant'],
            ['montant', 'Montant numérique'],
            ['devise', 'Code devise ISO'],
            ['commentaires', 'Commentaires métier'],
            ['source_pdf', 'PDF source (lien)'],
        ])

    # ── 5. Stack technique ──
    add_heading(doc, '5. Stack technique')
    add_heading(doc, 'Environnement de développement', level=2)
    add_bullet(doc, 'OS : Linux (Ubuntu 24.04 / WSL)')
    add_bullet(doc, 'IDE : VS Code avec extensions Python')
    add_bullet(doc, 'Gestion versions : Git + GitHub')
    add_bullet(doc, 'Déploiement cloud : Hugging Face Spaces')

    add_heading(doc, 'Dépendances principales', level=2)
    add_table(doc,
        ['Package', 'Version', 'Rôle'],
        [
            ['streamlit', '≥1.52', 'Framework web'],
            ['pandas', '≥2.3', 'Manipulation de données'],
            ['openpyxl', '≥3.1', 'Fichiers Excel'],
            ['PyMuPDF', '≥1.23', 'Extraction PDF (~50× plus rapide)'],
            ['pdfplumber', '≥0.11', 'Extraction PDF (fallback)'],
            ['pdfminer.six', '≥20251107', 'Moteur bas niveau PDF'],
            ['python-dateutil', '≥2.9', 'Parsing de dates'],
        ])

    # ── 6. Règles métier ──
    add_heading(doc, '6. Règles métier implémentées')
    add_para(doc, (
        "L'application intègre des règles métier spécifiques pour la classification et le traitement. "
        "Voir la Documentation Technique (sections 5.1 à 5.11) pour le détail exhaustif de chaque règle."
    ))

    add_heading(doc, '6.1 Exceptions MT202 sortants', level=2)
    add_bullet(doc, 'BEACCMCX091 dans F52A → Exception (annulable si receiver = SCBLGB2LXXX/CITIGB2LXXX)')
    add_bullet(doc, 'BC dans F21 → Exception salle des marchés (annulable)')

    add_heading(doc, '6.2 Exceptions MT202 entrants', level=2)
    add_bullet(doc, '323201 dans F58A → Exception')

    add_heading(doc, '6.3 Exceptions EUR', level=2)
    add_bullet(doc, 'T2PI (TARGET2 Paiement Instantané) → intérêts')
    add_bullet(doc, 'T2RM (TARGET2 Remboursement) → remboursement')
    add_bullet(doc, 'T2PL (TARGET2 Prélèvement) → placement')

    add_heading(doc, '6.4 Exceptions MT910', level=2)
    add_bullet(doc, 'Nivellement (NIVLT/5175)')
    add_bullet(doc, 'BEACCMCX091 dans F50A/F52A')
    add_bullet(doc, 'Forex (code dans liste bic_codes.xlsx/forex)')
    add_bullet(doc, 'Salle des marchés IBAN (FR7630001000640000005169558)')

    add_heading(doc, '6.5 Exception BANQUE DE FRANCE', level=2)
    add_bullet(doc, 'MT103 USD avec BANQUE DE FRANCE ou FW021083459 dans F53A/F54A/F57A')

    add_heading(doc, '6.6 Extraction commentaires', level=2)
    add_bullet(doc, 'MT202 : F72 « Texte descriptif: » (toutes occurrences, jointes par /)')
    add_bullet(doc, 'MT103 : F70 contenu complet')
    add_bullet(doc, "MT910 : F72 après « Info émetteur – destinataire » jusqu'à Block 5")
    add_bullet(doc, "Règle de non-écrasement : les commentaires d'exception ne sont pas remplacés")

    add_heading(doc, '6.7 Priorités extraction donneur d\'ordre', level=2)
    add_para(doc, 'MT202/MT910 entrants :', bold=True)
    add_bullet(doc, 'Priorité 1 : Code BIC dans F52A')
    add_bullet(doc, 'Priorité 2 : Code 4 chiffres dans F52D → Résolution via colonne Reglement')
    add_bullet(doc, 'Priorité 3 : Code BIC dans F52D')
    add_bullet(doc, 'Priorité 4 : Nom extrait de F52D')

    add_para(doc, 'MT103 entrants :', bold=True)
    add_bullet(doc, 'Priorité 1 : Code BIC dans F50K/F50F')
    add_bullet(doc, 'Priorité 2 : Nom après "Number: Numéro: 1/Details: Détails:" dans F50F')
    add_bullet(doc, 'Priorité 3 : Fallback nom extrait par parser')

    add_para(doc, 'MT103 sortants :', bold=True)
    add_bullet(doc, 'Priorité 1 : Codes Trésor (1001-6001) dans F50F')
    add_bullet(doc, 'Priorité 2 : Code BIC dans F52A')
    add_bullet(doc, 'Priorité 3 : Codes CCF 4 chiffres dans F50F')
    add_bullet(doc, 'Priorité 4 : Nom après "Details: Détails:" dans F50F')

    add_heading(doc, '6.8 Rapprochement MT950', level=2)
    add_bullet(doc, '4 critères cumulatifs : montant (±0.011), identification_code, date_valeur, référence')
    add_bullet(doc, 'Entrants : match exact sur ref_serv (après //) vs F20')
    add_bullet(doc, "Sortants : match préfixe sur ref_owner vs F20")
    add_bullet(doc, "Tout le pool de messages (y compris exceptions) est utilisé")

    # ── 7. Options de déploiement ──
    add_heading(doc, '7. Options de déploiement')

    add_heading(doc, '7.1 Serveur Interne (Recommandé)', level=2)
    add_para(doc, 'Configuration recommandée :', bold=True)
    add_bullet(doc, 'OS : Ubuntu 22.04+ ou Debian 11+')
    add_bullet(doc, 'Python : 3.10+')
    add_bullet(doc, 'Proxy : Nginx (reverse proxy + HTTPS)')
    add_bullet(doc, 'Firewall : Limiter accès au port 8501')
    add_bullet(doc, "Authentification : À implémenter selon besoins SI")
    add_bullet(doc, 'Stockage : Disque local persistant')
    add_bullet(doc, 'Logs : Rotation automatique via logrotate')
    add_bullet(doc, 'Monitoring : Supervision CPU/RAM/disque')

    add_para(doc, 'Commande de lancement :', bold=True)
    add_code_block(doc, 'streamlit run app.py --server.port 8501 --server.headless true')

    add_para(doc, 'Docker :', bold=True)
    add_code_block(doc, (
        "docker build -t pdf-swift-extractor .\n"
        "docker run -d --name pdf-extractor -p 8501:8501 \\\n"
        "  --restart unless-stopped pdf-swift-extractor"
    ))

    add_heading(doc, '7.2 Hugging Face Spaces (Demo/Test)', level=2)
    add_bullet(doc, 'URL : https://huggingface.co/spaces/Born237/SWIFT_EXTRACTOR')
    add_bullet(doc, 'Statut : ✓ Déployé')
    add_bullet(doc, '⚠️ Limitation : Pas de stockage persistant')
    add_bullet(doc, 'Usage : Démonstration, tests externes')

    add_heading(doc, '7.3 Local (Développement)', level=2)
    add_bullet(doc, 'URL : http://localhost:8501')
    add_bullet(doc, 'Usage : Tests, développement, formation utilisateurs')

    # ── 8. Caractéristiques techniques ──
    add_heading(doc, '8. Caractéristiques techniques')

    add_heading(doc, '8.1 Performance', level=2)
    add_bullet(doc, 'Traitement multi-messages : 120+ messages par PDF')
    add_bullet(doc, 'Extraction PDF : PyMuPDF (~50× plus rapide que pdfplumber)')
    add_bullet(doc, 'Cache BIC : Chargement unique du référentiel en mémoire')
    add_bullet(doc, '6 heuristiques de split en cascade pour robustesse')

    add_heading(doc, '8.2 Sécurité', level=2)
    add_bullet(doc, 'Aucune donnée stockée côté serveur — traitement en mémoire')
    add_bullet(doc, 'Pas d\'authentification native (à ajouter via reverse proxy)')
    add_bullet(doc, 'Fichiers traités en RAM uniquement, non persistés')

    add_heading(doc, '8.3 Fiabilité', level=2)
    add_bullet(doc, 'Extraction robuste : 6 heuristiques de split, 5 stratégies de détection type')
    add_bullet(doc, 'Fallback PDF : PyMuPDF → pdfplumber si échec')
    add_bullet(doc, 'Gestion des formats hétérogènes (labels EN + FR)')
    add_bullet(doc, 'Traçabilité complète via logs')

    add_heading(doc, '8.4 Maintenabilité', level=2)
    add_bullet(doc, 'Architecture modulaire (1 extracteur par type MT)')
    add_bullet(doc, 'Référentiel BIC externe modifiable sans recompilation')
    add_bullet(doc, 'Logs détaillés pour le débogage')

    # ── 9. Points d'attention SI ──
    add_heading(doc, '9. Points d\'attention pour le SI')

    add_heading(doc, '9.1 Sécurité réseau', level=2)
    add_bullet(doc, 'Port 8501 à restreindre au réseau interne')
    add_bullet(doc, 'HTTPS recommandé via reverse proxy Nginx')
    add_bullet(doc, 'Authentification à implémenter (LDAP, SSO…)')

    add_heading(doc, '9.2 Ressources serveur', level=2)
    add_bullet(doc, 'RAM : 2 Go minimum recommandé')
    add_bullet(doc, 'CPU : 2 cœurs minimum')
    add_bullet(doc, 'Disque : 1 Go pour l\'application + logs')

    add_heading(doc, '9.3 Sauvegarde et archivage', level=2)
    add_bullet(doc, 'Sauvegarder le fichier bic_codes.xlsx (référentiel métier)')
    add_bullet(doc, 'Archiver les logs applicatifs (rotation recommandée)')
    add_bullet(doc, 'Versionner le code source (Git)')

    add_heading(doc, '9.4 Maintenance applicative', level=2)
    add_bullet(doc, 'Mise à jour BIC : modifier bic_codes.xlsx + redémarrer')
    add_bullet(doc, 'Mise à jour code : remplacer fichiers + pip install -r requirements.txt + redémarrer')
    add_bullet(doc, 'Supervision : surveiller logs/app.log pour erreurs')

    add_heading(doc, '9.5 Scalabilité', level=2)
    add_bullet(doc, 'Application mono-instance (suffisant pour usage actuel)')
    add_bullet(doc, 'Si charge élevée : déployer plusieurs instances derrière un load balancer')
    add_bullet(doc, 'Docker facilite la réplication')

    # ── 10. Conclusion ──
    add_heading(doc, '10. Conclusion')
    add_para(doc, (
        "Le PDF SWIFT Extractor est une solution robuste et modulaire pour l'automatisation "
        "de l'extraction et de l'analyse des messages SWIFT. Son architecture en couches permet "
        "une maintenance facilitée et une évolution progressive des règles métier."
    ))

    add_heading(doc, 'Avantages clés', level=2)
    add_bullet(doc, 'Extraction automatisée : gain de temps considérable vs. traitement manuel')
    add_bullet(doc, 'Règles métier précises : 9+ exceptions différentes gérées automatiquement')
    add_bullet(doc, 'Traçabilité complète : logs, hyperliens, doublons détectés')
    add_bullet(doc, 'Rapprochement MT950 : 4 critères cumulatifs pour un matching fiable')
    add_bullet(doc, 'Déploiement flexible : local, Docker, ou cloud')

    add_heading(doc, 'Prochaines évolutions possibles', level=2)
    add_bullet(doc, 'Authentification utilisateurs intégrée')
    add_bullet(doc, 'Support de nouveaux types de messages SWIFT')
    add_bullet(doc, 'API REST pour intégration avec d\'autres systèmes')
    add_bullet(doc, 'Tableau de bord statistique (volumes, tendances)')

    # ── Save ──
    path = os.path.join(BASE, 'ARCHITECTURE_SWIFT_EXTRACTOR.docx')
    doc.save(path)
    print(f'✅ ARCHITECTURE créé : {path}')
    print(f'   → {len(doc.paragraphs)} paragraphes, {len(doc.tables)} tables')


# ─────────────────────────────────────────────────
if __name__ == '__main__':
    print('=== Mise à jour des documents Word ===')
    print()
    create_documentation_technique()
    print()
    create_architecture()
    print()
    print('✅ Terminé !')
