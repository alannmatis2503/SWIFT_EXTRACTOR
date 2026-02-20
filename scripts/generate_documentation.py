#!/usr/bin/env python3
"""
Génère le document .docx de documentation exhaustive pour le SWIFT Extractor.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

def add_heading(doc, text, level=1):
    doc.add_heading(text, level)

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = str(cell_data) if cell_data else ""
    
    return table

def main():
    doc = Document()
    
    # Title
    title = doc.add_heading('SWIFT EXTRACTOR - Documentation Technique', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f'Version: Janvier 2026\nDate de génération: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('')
    
    # =========================================================================
    # SECTION 1: INTRODUCTION
    # =========================================================================
    add_heading(doc, '1. Introduction', 1)
    add_paragraph(doc, """Ce document présente de manière exhaustive le fonctionnement de l'application SWIFT Extractor. Il détaille pour chaque type de message (MT103, MT202, MT900, MT910) et chaque direction (entrant/sortant) les champs SWIFT sources utilisés pour extraire les données des colonnes du classeur Excel final.""")
    
    add_paragraph(doc, """L'application traite les fichiers PDF contenant des messages SWIFT et génère un classeur Excel récapitulatif avec différentes feuilles selon les règles métier définies.""")
    
    # =========================================================================
    # SECTION 2: TYPES DE MESSAGES TRAITÉS
    # =========================================================================
    add_heading(doc, '2. Types de Messages SWIFT Traités', 1)
    
    add_heading(doc, '2.1 Messages Acceptés', 2)
    add_paragraph(doc, """Les types de messages suivants sont acceptés et traités par l'application:""")
    
    types_data = [
        ('fin.103', 'MT103', 'Virement client (single customer credit transfer)'),
        ('fin.202', 'MT202', 'Virement interbancaire (general financial institution transfer)'),
        ('fin.202.COV', 'MT202.COV', 'Virement interbancaire avec couverture'),
        ('fin.900', 'MT900', 'Confirmation de débit (debit confirmation)'),
        ('fin.910', 'MT910', 'Confirmation de crédit (credit confirmation)'),
    ]
    add_table(doc, ['Type Interne', 'Code SWIFT', 'Description'], types_data)
    
    add_heading(doc, '2.2 Messages Rejetés', 2)
    add_paragraph(doc, """Tout message SWIFT dont le type n'est pas 103, 202, 900 ou 910 est automatiquement rejeté et non inclus dans le classeur de sortie. Les messages MT900 ne sont traités que dans le mode "Analyse des transferts" (non disponible pour les modes incoming/outgoing classiques).""")
    
    # =========================================================================
    # SECTION 3: COLONNES DU CLASSEUR ET SOURCES
    # =========================================================================
    add_heading(doc, '3. Colonnes du Classeur Excel et Sources d\'Extraction', 1)
    
    add_paragraph(doc, """Le classeur Excel final contient les colonnes suivantes. Pour chaque colonne, les champs SWIFT source varient selon le type de message et la direction (entrant/sortant).""")
    
    # 3.1 Résumé des colonnes
    add_heading(doc, '3.1 Liste des Colonnes', 2)
    
    columns_data = [
        ('Code Banque', 'Code BIC de l\'établissement'),
        ('Date', 'Date de valeur du transfert'),
        ('Référence', 'Référence unique du message'),
        ('Type MT', 'Type de message SWIFT'),
        ('Pays', 'Code pays ISO 3 lettres'),
        ('Code Donneur d\'Ordre', 'Code BIC ou numérique du donneur d\'ordre'),
        ('Donneur d\'Ordre', 'Nom du donneur d\'ordre'),
        ('Bénéficiaire', 'Nom du bénéficiaire'),
        ('Correspondant', 'BIC du correspondant bancaire'),
        ('Montant', 'Montant du transfert'),
        ('Devise', 'Code devise ISO 3 lettres'),
        ('Commentaires', 'Remarques et motifs d\'exception'),
    ]
    add_table(doc, ['Colonne', 'Description'], columns_data)
    
    # =========================================================================
    # 3.2 MT103 ENTRANT (INCOMING)
    # =========================================================================
    add_heading(doc, '3.2 MT103 Entrant (Incoming)', 2)
    
    add_paragraph(doc, """Pour les messages MT103 en réception (incoming), les données sont extraites comme suit:""")
    
    mt103_in_data = [
        ('Code Banque', 'F52A', 'IdentifierCode / Code d\'identifiant'),
        ('Date', 'F32A', 'Date: section date du champ (format YYMMDD)'),
        ('Référence', 'F20', 'Référence du message'),
        ('Type MT', 'En-tête', 'Identifier: fin.103'),
        ('Pays', 'Mapping BIC', 'Obtenu depuis la colonne PAYS du fichier bic_codes.xlsx à partir du code BIC'),
        ('Code Donneur d\'Ordre', 'F50K / F50F', 'Code BIC à 8-11 caractères (positions 5-6 = code pays CEMAC)'),
        ('Donneur d\'Ordre', 'Priorité 1: BIC → nom via bic_codes.xlsx\nPriorité 2: F50F "Number: Numéro: 1/Details: Détails:"\nPriorité 3: Première ligne du NameAndAddress', 'Nom mappé ou extrait du champ'),
        ('Bénéficiaire', '(Vide)', 'Non renseigné pour les MT103 entrants'),
        ('Correspondant', 'En-tête Sender', 'BIC de l\'expéditeur du message'),
        ('Montant', 'F32A', 'Section Amount/Montant'),
        ('Devise', 'F32A', 'Section Currency/Devise'),
        ('Commentaires', 'F70', 'Contenu COMPLET du champ (Remittance Information)'),
    ]
    add_table(doc, ['Colonne', 'Champ SWIFT Source', 'Détail'], mt103_in_data)
    
    add_paragraph(doc, """IMPORTANT - Validation du code BIC pour les pays CEMAC:""", bold=True)
    add_paragraph(doc, """Seuls les codes BIC dont les positions 5-6 correspondent à un pays de la CEMAC sont acceptés:
- CM (Cameroun)
- CF (République Centrafricaine)
- CG (Congo)
- GA (Gabon)
- GQ (Guinée Équatoriale)
- TD (Tchad)
- FR (France - pour les banques françaises opérant en CEMAC)""")
    
    add_paragraph(doc, """Les mots français/anglais courants sont exclus de la détection BIC (faux positifs):""", bold=True)
    add_paragraph(doc, """GENERALE, FINANCES, BANQUE, RECETTE, PAIERIE, TRESOR, MINISTERE, CAISSE, COMPTABLE, DETAILS, NATIONALE""")
    
    # =========================================================================
    # 3.3 MT103 SORTANT (OUTGOING)
    # =========================================================================
    add_heading(doc, '3.3 MT103 Sortant (Outgoing)', 2)
    
    add_paragraph(doc, """Pour les messages MT103 en émission (outgoing), les données sont extraites selon les priorités suivantes:""")
    
    add_paragraph(doc, """Extraction du Donneur d'Ordre - Ordre de priorité:""", bold=True)
    add_paragraph(doc, """1. Codes Trésor (1001-6001) dans F50F → mappés via colonne "Reglement" de bic_codes.xlsx
2. Code BIC dans F52A → mappé via bic_codes.xlsx
3. Codes CCF 4 chiffres dans F50F → mappés via colonne "CCF" de bic_codes.xlsx
4. Fallback: nom extrait du NameAndAddress de F50F""")
    
    mt103_out_data = [
        ('Code Banque', 'F52A', 'IdentifierCode / Code d\'identifiant'),
        ('Date', 'F32A', 'Date: section date du champ'),
        ('Référence', 'F20', 'Référence du message'),
        ('Type MT', 'En-tête', 'Identifier: fin.103'),
        ('Pays', 'Mapping BIC', 'Depuis bic_codes.xlsx'),
        ('Code Donneur d\'Ordre', 'Priorité: Trésor → BIC → CCF', 'Voir priorités ci-dessus'),
        ('Donneur d\'Ordre', 'Priorité: Trésor → BIC → CCF → F50F name', 'Voir priorités ci-dessus'),
        ('Bénéficiaire', '(Vide)', 'Non renseigné pour les MT103 sortants'),
        ('Correspondant', 'En-tête Receiver', 'BIC du destinataire du message'),
        ('Montant', 'F32A', 'Section Amount/Montant'),
        ('Devise', 'F32A', 'Section Currency/Devise'),
        ('Commentaires', '(Vide)', 'Non extrait pour les sortants'),
    ]
    add_table(doc, ['Colonne', 'Champ SWIFT Source', 'Détail'], mt103_out_data)
    
    # =========================================================================
    # 3.4 MT202 ENTRANT (INCOMING)
    # =========================================================================
    add_heading(doc, '3.4 MT202 / MT202.COV Entrant (Incoming)', 2)
    
    mt202_in_data = [
        ('Code Banque', 'F52A / F52D', 'Priorité: F52A (BIC), sinon F52D (code 4 chiffres ou BIC)'),
        ('Date', 'F32A', 'ValueDate / Date de valeur'),
        ('Référence', 'F20', 'Transaction Reference'),
        ('Type MT', 'En-tête', 'Identifier: fin.202 ou fin.202.COV'),
        ('Pays', 'Mapping BIC', 'Depuis bic_codes.xlsx'),
        ('Code Donneur d\'Ordre', 'F52A / F52D', 'Code BIC ou code 4 chiffres (colonne Reglement)'),
        ('Donneur d\'Ordre', 'F52A / F52D', 'Nom mappé via bic_codes.xlsx'),
        ('Bénéficiaire', '(Vide)', 'Non renseigné pour les MT202 entrants'),
        ('Correspondant', 'En-tête Sender', 'BIC de l\'expéditeur'),
        ('Montant', 'F32A', 'Amount / Montant'),
        ('Devise', 'F32A', 'Currency / Devise'),
        ('Commentaires', 'F72', 'Contenu COMPLET du champ (Sender to Receiver Information)'),
    ]
    add_table(doc, ['Colonne', 'Champ SWIFT Source', 'Détail'], mt202_in_data)
    
    add_paragraph(doc, """Extraction du code depuis F52D:""", bold=True)
    add_paragraph(doc, """Si F52A n'est pas disponible ou ne contient pas de BIC valide, l'application cherche dans F52D:
1. D'abord un code 4 chiffres dans PartyIdentifier → mappé via colonne "Reglement" de bic_codes.xlsx
2. Sinon un code BIC → mappé via bic_codes.xlsx
3. Sinon le nom depuis NameAndAddress""")
    
    # =========================================================================
    # 3.5 MT202 SORTANT (OUTGOING)
    # =========================================================================
    add_heading(doc, '3.5 MT202 / MT202.COV Sortant (Outgoing)', 2)
    
    mt202_out_data = [
        ('Code Banque', 'F52A / F52D', 'Même logique que MT202 entrant'),
        ('Date', 'F32A', 'ValueDate / Date de valeur'),
        ('Référence', 'F20', 'Transaction Reference'),
        ('Type MT', 'En-tête', 'Identifier: fin.202 ou fin.202.COV'),
        ('Pays', 'Mapping BIC', 'Depuis bic_codes.xlsx'),
        ('Code Donneur d\'Ordre', 'F52A / F52D', 'Même logique que MT202 entrant'),
        ('Donneur d\'Ordre', 'F52A / F52D', 'Nom mappé via bic_codes.xlsx'),
        ('Bénéficiaire', 'F58A / F58D', 'Beneficiary Institution - Priorité: BIC → NameAndAddress'),
        ('Correspondant', 'En-tête Receiver', 'BIC du destinataire'),
        ('Montant', 'F32A', 'Amount / Montant'),
        ('Devise', 'F32A', 'Currency / Devise'),
        ('Commentaires', '(Vide)', 'Non extrait pour les sortants'),
    ]
    add_table(doc, ['Colonne', 'Champ SWIFT Source', 'Détail'], mt202_out_data)
    
    add_paragraph(doc, """Extraction du bénéficiaire depuis F58A/F58D:""", bold=True)
    add_paragraph(doc, """1. F58A (prioritaire): extraction du code BIC → mappé pour obtenir le nom
2. F58D (fallback): d'abord chercher un BIC avant NameAndAddress, sinon extraire le nom du NameAndAddress""")
    
    # =========================================================================
    # 3.6 MT910 ENTRANT (INCOMING)
    # =========================================================================
    add_heading(doc, '3.6 MT910 Entrant (Incoming)', 2)
    
    add_paragraph(doc, """Le MT910 est une confirmation de crédit. Le champ F52A représente à la fois le donneur d'ordre ET le bénéficiaire.""")
    
    mt910_in_data = [
        ('Code Banque', 'F52A', 'IdentifierCode'),
        ('Date', 'F32A', 'Value Date'),
        ('Référence', 'F20', 'Transaction Reference'),
        ('Type MT', 'En-tête', 'Identifier: fin.910'),
        ('Pays', 'Mapping BIC', 'FORCÉ depuis F52A BIC (priorité sur détection texte)'),
        ('Code Donneur d\'Ordre', 'F52A', 'Code BIC'),
        ('Donneur d\'Ordre', 'F52A', 'Nom mappé via bic_codes.xlsx'),
        ('Bénéficiaire', 'F52A', 'Identique au donneur d\'ordre'),
        ('Correspondant', 'En-tête Sender', 'BIC de l\'expéditeur'),
        ('Montant', 'F32A', 'Amount'),
        ('Devise', 'F32A', 'Currency'),
        ('Commentaires', '(Variable)', 'Si exception nivellement: "nivellement"'),
    ]
    add_table(doc, ['Colonne', 'Champ SWIFT Source', 'Détail'], mt910_in_data)
    
    # =========================================================================
    # 3.7 MT900 (ANALYSE DES TRANSFERTS)
    # =========================================================================
    add_heading(doc, '3.7 MT900 - Mode Analyse des Transferts', 2)
    
    add_paragraph(doc, """Le MT900 est une confirmation de débit. Il est traité uniquement dans le mode "Analyse des transferts exécutés".""")
    
    mt900_data = [
        ('Code Banque', 'F52A', 'IdentifierCode (si présent)'),
        ('Date', 'F32A', 'Value Date'),
        ('Référence', 'F20', 'Transaction Reference'),
        ('Référence d\'Origine', 'F21', 'Related Reference - utilisé pour matcher avec MT202/MT103'),
        ('Type MT', 'En-tête', 'Identifier: fin.900'),
        ('Pays', 'Matching MT202/MT103', 'Récupéré du message MT202/MT103 matché'),
        ('Code Donneur d\'Ordre', 'Matching MT202/MT103', 'Récupéré du message matché'),
        ('Donneur d\'Ordre', 'Matching MT202/MT103', 'Récupéré du message matché'),
        ('Bénéficiaire', 'Matching MT202/MT103', 'Récupéré du message matché'),
        ('Correspondant', 'Matching MT202/MT103', 'Récupéré du message matché'),
        ('Montant', 'F32A', 'Amount'),
        ('Devise', 'F32A', 'Currency'),
        ('Commentaires', '(Variable)', 'Si exception: "T2PL" ou "nivellement"'),
    ]
    add_table(doc, ['Colonne', 'Champ SWIFT Source', 'Détail'], mt900_data)
    
    add_paragraph(doc, """Le matching s'effectue en comparant la référence d'origine (F21) du MT900 avec la référence (F20) des messages MT202/MT103 du même fichier.""")
    
    # =========================================================================
    # SECTION 4: RÈGLES ET EXCEPTIONS
    # =========================================================================
    add_heading(doc, '4. Règles d\'Exclusion et Exceptions', 1)
    
    # 4.1 Règles d'exclusion
    add_heading(doc, '4.1 Règle 1: Messages MT103 USD - Exclusion BANQUE DE FRANCE', 2)
    
    add_paragraph(doc, """Champ d'application:""", bold=True)
    add_paragraph(doc, """Cette règle s'applique UNIQUEMENT aux messages MT103 en devise USD.""")
    
    add_paragraph(doc, """Condition d'exclusion:""", bold=True)
    add_paragraph(doc, """Un message MT103 en USD est EXCLU du classeur si l'un des champs F53A, F54A ou F57A contient:
- "BANQUE DE FRANCE"
- "FW021083459"

Ces champs correspondent aux correspondants bancaires intermédiaires.""")
    
    add_paragraph(doc, """Impact:""", bold=True)
    add_paragraph(doc, """Le message n'apparaît dans AUCUNE feuille du classeur (ni summary, ni exceptions).""")
    
    # 4.2 Messages BEACCMCX091
    add_heading(doc, '4.2 Règle 2: Messages MT910 - BEACCMCX091', 2)
    
    add_paragraph(doc, """Champ d'application:""", bold=True)
    add_paragraph(doc, """Messages MT910 (confirmation de crédit) uniquement.""")
    
    add_paragraph(doc, """Condition:""", bold=True)
    add_paragraph(doc, """Le code BEACCMCX091 est détecté dans F50A ou F52A (après "IdentifierCode: Code d'identifiant:").""")
    
    add_paragraph(doc, """Impact:""", bold=True)
    add_paragraph(doc, """Le message est placé dans une feuille séparée "BEACCMCX091" au lieu de la feuille "summary".""")
    
    # 4.3 Exception 323201
    add_heading(doc, '4.3 Règle 3: Messages MT202 Entrants - Exception 323201', 2)
    
    add_paragraph(doc, """Champ d'application:""", bold=True)
    add_paragraph(doc, """Messages MT202 entrants (incoming) uniquement.""")
    
    add_paragraph(doc, """Condition:""", bold=True)
    add_paragraph(doc, """Le champ F58A (Beneficiary Institution) contient la séquence "323201".""")
    
    add_paragraph(doc, """Impact:""", bold=True)
    add_paragraph(doc, """Le message est placé dans une feuille séparée "Exceptions_323201".""")
    
    # 4.4 Exceptions EUR
    add_heading(doc, '4.4 Règle 4: Exceptions EUR - T2PI, T2RM, T2PL', 2)
    
    add_paragraph(doc, """Champ d'application:""", bold=True)
    add_paragraph(doc, """Messages en devise EUR uniquement.""")
    
    add_paragraph(doc, """Conditions et commentaires:""", bold=True)
    
    eur_rules_data = [
        ('T2PI', 'Référence (F20)', 'Entrant ET Sortant', 'intérêts'),
        ('T2RM', 'Référence (F20)', 'Entrant uniquement', 'remboursement'),
        ('T2PL', 'Référence (F20)', 'Sortant uniquement', 'placement'),
    ]
    add_table(doc, ['Code', 'Champ Source', 'Direction', 'Commentaire Ajouté'], eur_rules_data)
    
    add_paragraph(doc, """Impact:""", bold=True)
    add_paragraph(doc, """Le message est placé dans une feuille séparée "Autres_Exceptions".""")
    
    # 4.5 Exception Nivellement MT910
    add_heading(doc, '4.5 Règle 5: Exceptions Nivellement - MT910', 2)
    
    add_paragraph(doc, """Champ d'application:""", bold=True)
    add_paragraph(doc, """Messages MT910 uniquement.""")
    
    add_paragraph(doc, """Conditions:""", bold=True)
    
    niv_rules_data = [
        ('Entrant', 'F25P contient "5175" OU référence contient "NIVLT"'),
        ('Sortant', 'F53B ou F58A contient "5175"'),
    ]
    add_table(doc, ['Direction', 'Condition'], niv_rules_data)
    
    add_paragraph(doc, """Impact:""", bold=True)
    add_paragraph(doc, """Le message est placé dans "Autres_Exceptions" avec commentaire "nivellement".""")
    
    # 4.6 Exception MT900
    add_heading(doc, '4.6 Règle 6: Exceptions MT900 (Analyse des transferts)', 2)
    
    add_paragraph(doc, """Conditions:""", bold=True)
    
    mt900_rules_data = [
        ('F20 (référence)', 'Contient "T2PL"', 'T2PL'),
        ('F21 (référence d\'origine)', 'Contient "NIVLT" ou "NIVELLEMENT"', 'nivellement'),
    ]
    add_table(doc, ['Champ', 'Condition', 'Commentaire Ajouté'], mt900_rules_data)
    
    add_paragraph(doc, """Impact:""", bold=True)
    add_paragraph(doc, """Le message MT900 n'est PAS matché avec les MT202/MT103 et va dans une feuille d'exceptions séparée.""")
    
    # =========================================================================
    # SECTION 5: STRUCTURE DU CLASSEUR EXCEL
    # =========================================================================
    add_heading(doc, '5. Structure du Classeur Excel Généré', 1)
    
    add_paragraph(doc, """Le classeur Excel généré contient les feuilles suivantes, dans l'ordre:""")
    
    sheets_data = [
        ('summary', 'Obligatoire', 'Récapitulatif de tous les messages valides'),
        ('BEACCMCX091', 'Conditionnelle', 'Messages MT910 avec code BEACCMCX091'),
        ('Exceptions_323201', 'Conditionnelle', 'Messages MT202 entrants avec 323201 dans F58A'),
        ('Autres_Exceptions', 'Conditionnelle', 'Exceptions EUR (T2PI/T2RM/T2PL) et nivellement'),
        ('Doublons_potentiels', 'Conditionnelle', 'Messages avec référence identique'),
        ('[Code Pays]', 'Multiple', 'Une feuille par pays (ex: CM, GA, etc.)'),
        ('[Nom Message]', 'Multiple', 'Une feuille détaillée par message individuel'),
    ]
    add_table(doc, ['Nom de Feuille', 'Présence', 'Contenu'], sheets_data)
    
    # =========================================================================
    # SECTION 6: FICHIER BIC_CODES.XLSX
    # =========================================================================
    add_heading(doc, '6. Structure du Fichier bic_codes.xlsx', 1)
    
    add_paragraph(doc, """Le fichier bic_codes.xlsx est essentiel pour le mapping des codes vers les noms d'établissements. Il doit contenir les colonnes suivantes:""")
    
    bic_cols_data = [
        ('Code BIC / BIC', 'Obligatoire', 'Code BIC à 8-11 caractères'),
        ('Nom / Name', 'Obligatoire', 'Nom de l\'établissement bancaire'),
        ('Pays / Country', 'Recommandée', 'Code pays ISO 2 lettres'),
        ('Reglement', 'Optionnelle', 'Codes 4 chiffres (ex: 1001, 6001) pour le Trésor'),
        ('CCF', 'Optionnelle', 'Codes CCF 4 chiffres'),
    ]
    add_table(doc, ['Colonne', 'Obligatoire', 'Description'], bic_cols_data)
    
    add_paragraph(doc, """Emplacements recherchés pour bic_codes.xlsx:""", bold=True)
    add_paragraph(doc, """1. Variable d'environnement PDF_SWIFT_DATA_DIR
2. Dossier ProgramData (Windows)
3. Dossier utilisateur local
4. data/bic_codes.xlsx (relatif à l'application)""")
    
    # =========================================================================
    # SECTION 7: RÉSUMÉ DES DIRECTIVES IMPLÉMENTÉES
    # =========================================================================
    add_heading(doc, '7. Résumé des Directives Implémentées', 1)
    
    directives_data = [
        ('1', 'Types acceptés: 103, 202, 900, 910', '✅ Implémenté'),
        ('2', 'Exclusion MT103 USD avec BANQUE DE FRANCE', '✅ Implémenté'),
        ('3', 'Séparation BEACCMCX091 en feuille dédiée', '✅ Implémenté'),
        ('4', 'Exception 323201 pour MT202 entrants', '✅ Implémenté'),
        ('5', 'Exceptions EUR: T2PI (entrant/sortant), T2RM (entrant), T2PL (sortant)', '✅ Implémenté'),
        ('6', 'Exceptions nivellement MT910', '✅ Implémenté'),
        ('7', 'Exceptions MT900: T2PL dans F20, NIVLT/NIVELLEMENT dans F21', '✅ Implémenté'),
        ('8', 'Validation BIC: pays CEMAC uniquement (CM, CF, CG, GA, GQ, TD)', '✅ Implémenté'),
        ('9', 'Exclusion mots français des détections BIC', '✅ Implémenté'),
        ('10', 'Extraction donneur MT103 entrant: BIC → F50F Details → Fallback', '✅ Implémenté'),
        ('11', 'Extraction donneur MT103 sortant: Trésor → BIC → CCF → Fallback', '✅ Implémenté'),
        ('12', 'Extraction bénéficiaire MT202 sortant: F58A → F58D', '✅ Implémenté'),
        ('13', 'Commentaires: F72 (MT202 entrant), F70 (MT103 entrant) - contenu complet', '✅ Implémenté'),
        ('14', 'Bénéficiaire vide pour MT103 entrant et sortant', '✅ Implémenté'),
        ('15', 'Bénéficiaire vide pour MT202 entrant', '✅ Implémenté'),
        ('16', 'MT910: Pays forcé depuis BIC (priorité sur détection texte)', '✅ Implémenté'),
    ]
    add_table(doc, ['N°', 'Directive', 'Statut'], directives_data)
    
    # =========================================================================
    # FOOTER
    # =========================================================================
    doc.add_page_break()
    add_heading(doc, 'Annexe: Glossaire des Champs SWIFT', 1)
    
    glossary_data = [
        ('F20', 'Transaction Reference Number', 'Référence unique du message'),
        ('F21', 'Related Reference', 'Référence d\'un message lié (pour MT900)'),
        ('F25P', 'Account Identification', 'Identification du compte'),
        ('F32A', 'Value Date/Currency/Amount', 'Date valeur, devise et montant'),
        ('F50', 'Ordering Customer', 'Client donneur d\'ordre (variations: F50K, F50F)'),
        ('F52A', 'Ordering Institution', 'Institution donneur d\'ordre (format BIC)'),
        ('F52D', 'Ordering Institution', 'Institution donneur d\'ordre (format nom/adresse)'),
        ('F53A', 'Sender\'s Correspondent', 'Correspondant de l\'émetteur'),
        ('F53B', 'Sender\'s Correspondent', 'Correspondant de l\'émetteur (format alternatif)'),
        ('F54A', 'Receiver\'s Correspondent', 'Correspondant du destinataire'),
        ('F57A', 'Account With Institution', 'Institution tenant le compte'),
        ('F58A', 'Beneficiary Institution', 'Institution bénéficiaire (format BIC)'),
        ('F58D', 'Beneficiary Institution', 'Institution bénéficiaire (format nom/adresse)'),
        ('F59', 'Beneficiary Customer', 'Client bénéficiaire'),
        ('F70', 'Remittance Information', 'Informations sur le versement'),
        ('F72', 'Sender to Receiver Information', 'Information de l\'émetteur au destinataire'),
    ]
    add_table(doc, ['Code Champ', 'Nom SWIFT', 'Description'], glossary_data)
    
    # Save the document
    output_path = '/home/alannmatis/pdf-extractor/DOCUMENTATION_TECHNIQUE_SWIFT_EXTRACTOR.docx'
    doc.save(output_path)
    print(f"✅ Document généré: {output_path}")
    return output_path

if __name__ == '__main__':
    main()
