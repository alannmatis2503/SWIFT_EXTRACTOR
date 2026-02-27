#!/usr/bin/env python3
"""Update section 5.8 (Exception Forex) in the .docx documentation."""
import sys
from docx import Document
from docx.table import Table
from docx.oxml.ns import qn
from copy import deepcopy

def update_docx(docx_path):
    doc = Document(docx_path)
    body = doc.element.body
    elements = list(body)
    
    # Find section 5.8 heading
    heading_idx = None
    for idx, el in enumerate(elements):
        if el.tag.endswith('}p'):
            all_text = ''.join(node.text or '' for node in el.iter())
            if '5.8' in all_text and 'Forex' in all_text:
                heading_idx = idx
                break
    
    if heading_idx is None:
        print("ERROR: Section 5.8 not found")
        return False
    
    # Find the table right after (element heading_idx + 1)
    tbl_idx = heading_idx + 1
    if not elements[tbl_idx].tag.endswith('}tbl'):
        print(f"ERROR: Expected table at index {tbl_idx}")
        return False
    
    tbl = Table(elements[tbl_idx], doc)
    
    # Update row 3 (Condition) to show both rules
    row3 = tbl.rows[3]
    row3.cells[0].paragraphs[0].clear()  
    row3.cells[0].paragraphs[0].text = "Condition (règle 1)"
    row3.cells[1].paragraphs[0].clear()
    row3.cells[1].paragraphs[0].text = "Les 8 premiers caractères du code_donneur_dordre figurent dans la liste des codes forex"
    
    # Add a new row for rule 2
    # Clone row 3 to get formatting, then modify
    new_row_xml = deepcopy(tbl.rows[3]._tr)
    # Insert after row 3
    tbl.rows[3]._tr.addnext(new_row_xml)
    
    # Now the table has 7 rows. Access the new row (index 4)
    new_row = tbl.rows[4]
    new_row.cells[0].paragraphs[0].clear()
    new_row.cells[0].paragraphs[0].text = "Condition (règle 2)"
    new_row.cells[1].paragraphs[0].clear()
    new_row.cells[1].paragraphs[0].text = (
        "Si la règle 1 échoue ET que le correspondant commence par CITIUS33 "
        "ET que la devise est USD : on cherche les codes forex dans le champ F50A "
        "du message (recherche de sous-chaîne)"
    )
    
    # Add a note paragraph after the table
    # Find position right after the table in the body
    next_idx = tbl_idx + 1
    
    # Create a new paragraph
    from docx.oxml import OxmlElement
    new_para = OxmlElement('w:p')
    
    # Add paragraph properties for Normal style
    pPr = OxmlElement('w:pPr')
    new_para.append(pPr)
    
    # Add run with text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # italics
    i_el = OxmlElement('w:i')
    rPr.append(i_el)
    run.append(rPr)
    
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = (
        "Note : La règle 2 (CITIUS33 + USD) permet de détecter les opérations forex "
        "lorsque le code donneur d'ordre extrait de F52A n'est pas dans la liste forex, "
        "mais que le champ F50A contient un code forex. Le donneur d'ordre issu de F52A "
        "est conservé tel quel."
    )
    run.append(t)
    new_para.append(run)
    
    # Insert the note paragraph after the table
    elements[tbl_idx].addnext(new_para)
    
    doc.save(docx_path)
    print(f"OK: Section 5.8 mise à jour dans {docx_path}")
    return True


if __name__ == "__main__":
    # Update both copies
    paths = [
        "livraison_dsi/DOCUMENTATION_TECHNIQUE_SWIFT_EXTRACTOR.docx",
        "DOCUMENTATION_TECHNIQUE_SWIFT_EXTRACTOR.docx",
    ]
    import os
    for p in paths:
        if os.path.exists(p):
            update_docx(p)
        else:
            print(f"SKIP: {p} not found")
